import pandas as pd
from datetime import datetime
from typing import Dict, Any, Optional, Union
import os
from pathlib import Path
from weasyprint import HTML, CSS
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Mm
import tempfile
import traceback

class DocumentGenerator:
    """Generates various billing documents from processed Excel data"""
    
    def __init__(self, data: Dict[str, Any]):
        self.data = data
        self.title_data = data.get('title_data', {})
        self.work_order_data = data.get('work_order_data', pd.DataFrame())
        self.bill_quantity_data = data.get('bill_quantity_data', pd.DataFrame())
        self.extra_items_data = data.get('extra_items_data', pd.DataFrame())
    
    def _safe_float(self, value):
        """Safely convert value to float, return 0 if conversion fails"""
        try:
            if pd.isna(value) or value == '' or value is None:
                return 0.0
            return float(value)
        except (ValueError, TypeError):
            return 0.0
    
    def _find_column(self, df, possible_names):
        """Find column by trying multiple possible names"""
        for name in possible_names:
            if name in df.columns:
                return name
        return None

    def _generate_pdf(self, html_content: str, output_path: str = None, landscape: bool = False) -> Optional[bytes]:
        """
        Generate PDF from HTML with consistent 10mm margins and proper table layout
        
        Args:
            html_content: HTML content to convert
            output_path: Optional path to save PDF (if None, returns bytes)
            landscape: Whether to use landscape orientation
            
        Returns:
            PDF content as bytes if output_path is None, else None
        """
        try:
            # Base CSS with proper margins and table layout
            css_text = f"""
            @page {{
                size: A4 {'landscape' if landscape else 'portrait'};
                margin: 10mm;
                font-family: "Times New Roman", serif;
            }}
            body {{
                font-size: 10pt;
                line-height: 1.3;
                margin: 0;
                padding: 0;
            }}
            table {{
                width: 100%;
                table-layout: fixed;
                border-collapse: collapse;
                margin: 10px 0;
            }}
            th, td {{
                border: 1px solid #000;
                padding: 4pt 6pt;
                overflow: hidden;
                word-wrap: break-word;
            }}
            .header {{
                text-align: center;
                margin-bottom: 10mm;
            }}
            .title {{
                font-size: 12pt;
                font-weight: bold;
            }}
            .subtitle {{
                font-size: 10pt;
                margin: 5px 0;
            }}
            .amount {{
                text-align: right;
            }}
            """
            
            # Create HTML object
            html = HTML(string=html_content)
            
            # Generate PDF
            if output_path:
                html.write_pdf(
                    output_path,
                    stylesheets=[CSS(string=css_text)],
                    optimize_size=('fonts', 'images')
                )
                return None
            else:
                return html.write_pdf(stylesheets=[CSS(string=css_text)])
                
        except Exception as e:
            error_msg = f"Error in _generate_pdf: {str(e)}\n\n{traceback.format_exc()}"
            print(error_msg)
            raise
    
    def _generate_docx(self, html_content: str, output_path: str = None) -> Optional[bytes]:
        """
        Generate DOCX from HTML with consistent 10mm margins
        
        Args:
            html_content: HTML content to convert
            output_path: Optional path to save DOCX (if None, returns bytes)
            
        Returns:
            DOCX content as bytes if output_path is None, else None
        """
        doc = Document()
        
        # Set page margins to 10mm (converted to twips: 1mm = 56.7 twips)
        section = doc.sections[0]
        for margin in [section.top_margin, section.bottom_margin, 
                      section.left_margin, section.right_margin]:
            margin.mm = 10
        
        # Parse HTML and convert to DOCX
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Add title if present
        title = soup.find('title')
        if title:
            doc.add_heading(title.text, level=1)
        
        # Add tables
        for table in soup.find_all('table'):
            doc_table = doc.add_table(rows=0, cols=0)
            
            # Add header row if exists
            header_row = table.find('tr')
            if header_row:
                # Get column count from header
                col_count = len(header_row.find_all(['th', 'td']))
                doc_table.add_row()
                
                # Add header cells
                for i, cell in enumerate(header_row.find_all(['th', 'td'])):
                    if i >= col_count:
                        break
                    if i >= len(doc_table.rows[0].cells):
                        doc_table.add_column(0)
                    doc_table.rows[0].cells[i].text = cell.get_text(strip=True)
            
            # Add data rows
            for row in table.find_all('tr')[1:]:  # Skip header if it exists
                cells = row.find_all(['td', 'th'])
                if not cells:
                    continue
                    
                doc_row = doc_table.add_row()
                for i, cell in enumerate(cells):
                    if i >= col_count:  # Don't exceed column count from header
                        break
                    if i >= len(doc_row.cells):
                        doc_table.add_column(0)
                    doc_row.cells[i].text = cell.get_text(strip=True)
        
        if output_path:
            doc.save(output_path)
            return None
        else:
            temp_path = os.path.join(tempfile.gettempdir(), 'temp_docx.docx')
            doc.save(temp_path)
            with open(temp_path, 'rb') as f:
                return f.read()

    def generate_all_documents(self, output_format: str = 'html') -> Dict[str, Union[str, bytes]]:
        """
        Generate all required documents in the specified format
        
        Args:
            output_format: Output format ('html', 'pdf', or 'docx')
            
        Returns:
            Dictionary containing all generated documents
        """
        documents = {}
        
        # Generate individual documents
        html_docs = {
            'First Page Summary': self._generate_first_page(),
            'Deviation Statement': self._generate_deviation_statement(),
            'Final Bill Scrutiny Sheet': self._generate_final_bill_scrutiny(),
            'Extra Items Statement': self._generate_extra_items_statement(),
            'Certificate II': self._generate_certificate_ii(),
            'Certificate III': self._generate_certificate_iii()
        }
        
        # Convert to requested format
        if output_format.lower() == 'html':
            return html_docs
        
        for name, html_content in html_docs.items():
            if output_format.lower() == 'pdf':
                is_landscape = 'Deviation Statement' in name  # Example: Set landscape for specific documents
                documents[f"{name}.pdf"] = self._generate_pdf(html_content, landscape=is_landscape)
            elif output_format.lower() == 'docx':
                documents[f"{name}.docx"] = self._generate_docx(html_content)
        
        return documents
    
    def create_pdf_documents(self, documents: Dict[str, str]) -> Dict[str, bytes]:
        """
        Convert HTML documents to PDF format
        
        Args:
            documents: Dictionary of HTML documents
            
        Returns:
            Dictionary of PDF documents as bytes
        """
        pdf_files = {}
        
        for doc_name, html_content in documents.items():
            is_landscape = 'Deviation Statement' in doc_name  # Set landscape for specific documents
            pdf_content = self._generate_pdf(html_content, landscape=is_landscape)
            pdf_files[f"{doc_name}.pdf"] = pdf_content
        
        return pdf_files
    
    def _generate_first_page(self) -> str:
        """Generate First Page Summary document"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>First Page Summary</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 10mm; }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 18px; font-weight: bold; }}
                .subtitle {{ font-size: 14px; margin: 5px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ border: 1px solid #000; padding: 5px; text-align: left; }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">🏛️ Infrastructure Billing System</div>
                <div class="subtitle">First Page Summary</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <h3>Project Information</h3>
            <table>
                <tr><td><strong>Project Name:</strong></td><td>{self.title_data.get('Project Name', 'N/A')}</td></tr>
                <tr><td><strong>Contract No:</strong></td><td>{self.title_data.get('Contract No', 'N/A')}</td></tr>
                <tr><td><strong>Work Order No:</strong></td><td>{self.title_data.get('Work Order No', 'N/A')}</td></tr>
            </table>
            
            <h3>Work Items Summary</h3>
            <table>
                <thead>
                    <tr>
                        <th width="11.1mm">Item No.</th>
                        <th width="74.2mm">Item of Work supplies</th>
                        <th width="11.7mm">Unit</th>
                        <th width="16mm">Quantity executed since last certificate</th>
                        <th width="16mm">Quantity executed upto date</th>
                        <th width="15.3mm">Rate</th>
                        <th width="22.7mm">Amount upto date</th>
                        <th width="17.6mm">Amount Since previous bill</th>
                        <th width="13.9mm">Remark</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        # Add work order items
        for index, row in self.work_order_data.iterrows():
            # Safely convert numeric values
            quantity = self._safe_float(row.get('Quantity', 0))
            rate = self._safe_float(row.get('Rate', 0))
            amount = quantity * rate
            
            html_content += f"""
                    <tr>
                        <td>{row.get('Item', '')}</td>
                        <td>{row.get('Description', '')}</td>
                        <td>{row.get('Unit', '')}</td>
                        <td class="amount">{quantity:.2f}</td>
                        <td class="amount">{quantity:.2f}</td>
                        <td class="amount">{rate:.2f}</td>
                        <td class="amount">{amount:.2f}</td>
                        <td class="amount">{amount:.2f}</td>
                        <td></td>
                    </tr>
            """
        
        # Calculate totals
        total_amount = 0
        for _, row in self.work_order_data.iterrows():
            quantity = self._safe_float(row.get('Quantity', 0))
            rate = self._safe_float(row.get('Rate', 0))
            total_amount += quantity * rate
        
        html_content += f"""
                    <tr style="font-weight: bold;">
                        <td colspan="6">TOTAL</td>
                        <td class="amount">{total_amount:.2f}</td>
                        <td class="amount">{total_amount:.2f}</td>
                        <td></td>
                    </tr>
                </tbody>
            </table>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_deviation_statement(self) -> str:
        """Generate Deviation Statement document"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Deviation Statement</title>
            <style>
                @page {{ size: A4 landscape; margin: 10mm; }}
                body {{ font-family: Arial, sans-serif; margin: 0; }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 18px; font-weight: bold; }}
                .subtitle {{ font-size: 14px; margin: 5px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ border: 1px solid #000; padding: 3px; text-align: left; font-size: 10px; }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">🏛️ Infrastructure Billing System</div>
                <div class="subtitle">Deviation Statement</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <table>
                <thead>
                    <tr>
                        <th width="6mm">ITEM No.</th>
                        <th width="95mm">Description</th>
                        <th width="10mm">Unit</th>
                        <th width="10mm">Qty as per Work Order</th>
                        <th width="12mm">Rate</th>
                        <th width="12mm">Amt as per Work Order Rs.</th>
                        <th width="12mm">Qty Executed</th>
                        <th width="12mm">Amt as per Executed Rs.</th>
                        <th width="12mm">Excess Qty</th>
                        <th width="12mm">Excess Amt Rs.</th>
                        <th width="12mm">Saving Qty</th>
                        <th width="12mm">Saving Amt Rs.</th>
                        <th width="40mm">REMARKS/ REASON.</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        # Compare work order with bill quantity data
        for index, wo_row in self.work_order_data.iterrows():
            # Find corresponding bill quantity row
            bq_row = None
            if not self.bill_quantity_data.empty:
                matching_rows = self.bill_quantity_data[
                    self.bill_quantity_data['Item'] == wo_row.get('Item')
                ]
                if not matching_rows.empty:
                    bq_row = matching_rows.iloc[0]
            
            wo_qty = self._safe_float(wo_row.get('Quantity', 0))
            wo_rate = self._safe_float(wo_row.get('Rate', 0))
            wo_amount = wo_qty * wo_rate
            
            bq_qty = self._safe_float(bq_row.get('Quantity', 0)) if bq_row is not None else 0
            bq_amount = bq_qty * wo_rate
            
            excess_qty = max(0, bq_qty - wo_qty)
            excess_amt = excess_qty * wo_rate
            saving_qty = max(0, wo_qty - bq_qty)
            saving_amt = saving_qty * wo_rate
            
            html_content += f"""
                    <tr>
                        <td>{wo_row.get('Item', '')}</td>
                        <td>{wo_row.get('Description', '')}</td>
                        <td>{wo_row.get('Unit', '')}</td>
                        <td class="amount">{wo_qty:.2f}</td>
                        <td class="amount">{wo_rate:.2f}</td>
                        <td class="amount">{wo_amount:.2f}</td>
                        <td class="amount">{bq_qty:.2f}</td>
                        <td class="amount">{bq_amount:.2f}</td>
                        <td class="amount">{excess_qty:.2f}</td>
                        <td class="amount">{excess_amt:.2f}</td>
                        <td class="amount">{saving_qty:.2f}</td>
                        <td class="amount">{saving_amt:.2f}</td>
                        <td></td>
                    </tr>
            """
        
        html_content += """
                </tbody>
            </table>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_final_bill_scrutiny(self) -> str:
        """Generate Final Bill Scrutiny Sheet"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Final Bill Scrutiny Sheet</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 10mm; }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 18px; font-weight: bold; }}
                .subtitle {{ font-size: 14px; margin: 5px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ border: 1px solid #000; padding: 5px; text-align: left; }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">🏛️ Infrastructure Billing System</div>
                <div class="subtitle">Final Bill Scrutiny Sheet</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <h3>Bill Summary</h3>
            <table>
                <thead>
                    <tr>
                        <th>Item No.</th>
                        <th>Description</th>
                        <th>Unit</th>
                        <th>Quantity</th>
                        <th>Rate</th>
                        <th>Amount</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        total_amount = 0
        for index, row in self.bill_quantity_data.iterrows():
            quantity = self._safe_float(row.get('Quantity', 0))
            rate = self._safe_float(row.get('Rate', 0))
            amount = quantity * rate
            total_amount += amount
            
            html_content += f"""
                    <tr>
                        <td>{row.get('Item', '')}</td>
                        <td>{row.get('Description', '')}</td>
                        <td>{row.get('Unit', '')}</td>
                        <td class="amount">{quantity:.2f}</td>
                        <td class="amount">{rate:.2f}</td>
                        <td class="amount">{amount:.2f}</td>
                    </tr>
            """
        
        html_content += f"""
                    <tr style="font-weight: bold;">
                        <td colspan="5">TOTAL</td>
                        <td class="amount">{total_amount:.0f}</td>
                    </tr>
                </tbody>
            </table>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_extra_items_statement(self) -> str:
        """Generate Extra Items Statement"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Extra Items Statement</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 10mm; }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 18px; font-weight: bold; }}
                .subtitle {{ font-size: 14px; margin: 5px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ border: 1px solid #000; padding: 5px; text-align: left; }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">🏛️ Infrastructure Billing System</div>
                <div class="subtitle">Extra Items Statement</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
        """
        
        if not self.extra_items_data.empty:
            html_content += """
            <h3>Extra Items</h3>
            <table>
                <thead>
                    <tr>
                        <th>Item No.</th>
                        <th>Description</th>
                        <th>Unit</th>
                        <th>Quantity</th>
                        <th>Rate</th>
                        <th>Amount</th>
                    </tr>
                </thead>
                <tbody>
            """
            
            total_amount = 0
            for index, row in self.extra_items_data.iterrows():
                amount = row.get('Quantity', 0) * row.get('Rate', 0)
                total_amount += amount
                
                html_content += f"""
                        <tr>
                            <td>{row.get('Item No', '')}</td>
                            <td>{row.get('Description', '')}</td>
                            <td>{row.get('Unit', '')}</td>
                            <td class="amount">{row.get('Quantity', 0):.0f}</td>
                            <td class="amount">{row.get('Rate', 0):.0f}</td>
                            <td class="amount">{amount:.0f}</td>
                        </tr>
                """
            
            html_content += f"""
                        <tr style="font-weight: bold;">
                            <td colspan="5">TOTAL</td>
                            <td class="amount">{total_amount:.0f}</td>
                        </tr>
                    </tbody>
                </table>
            """
        else:
            html_content += "<p>No extra items found in the provided data.</p>"
        
        html_content += """
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_certificate_ii(self) -> str:
        """Generate Certificate II"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Certificate II</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 10mm; }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 18px; font-weight: bold; }}
                .subtitle {{ font-size: 14px; margin: 5px 0; }}
                .content {{ margin: 20px 0; line-height: 1.6; }}
                .signature {{ margin-top: 50px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">🏛️ Infrastructure Billing System</div>
                <div class="subtitle">Certificate II</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <div class="content">
                <p>This is to certify that the work described in the bill has been executed according to the specifications and is complete in all respects.</p>
                
                <p><strong>Project Details:</strong></p>
                <ul>
                    <li>Project Name: {self.title_data.get('Project Name', 'N/A')}</li>
                    <li>Contract No: {self.title_data.get('Contract No', 'N/A')}</li>
                    <li>Work Order No: {self.title_data.get('Work Order No', 'N/A')}</li>
                </ul>
                
                <p>The work has been executed in accordance with the contract terms and conditions.</p>
            </div>
            
            <div class="signature">
                <p>_________________________</p>
                <p>Engineer-in-Charge</p>
                <p>Date: {current_date}</p>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_certificate_iii(self) -> str:
        """Generate Certificate III"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Certificate III</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 10mm; }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 18px; font-weight: bold; }}
                .subtitle {{ font-size: 14px; margin: 5px 0; }}
                .content {{ margin: 20px 0; line-height: 1.6; }}
                .signature {{ margin-top: 50px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">🏛️ Infrastructure Billing System</div>
                <div class="subtitle">Certificate III</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <div class="content">
                <p>This is to certify that the rates charged in the bill are in accordance with the contract and approved rate schedule.</p>
                
                <p><strong>Project Details:</strong></p>
                <ul>
                    <li>Project Name: {self.title_data.get('Project Name', 'N/A')}</li>
                    <li>Contract No: {self.title_data.get('Contract No', 'N/A')}</li>
                    <li>Work Order No: {self.title_data.get('Work Order No', 'N/A')}</li>
                </ul>
                
                <p>All rates and calculations have been verified and are correct.</p>
            </div>
            
            <div class="signature">
                <p>_________________________</p>
                <p>Accounts Officer</p>
                <p>Date: {current_date}</p>
            </div>
        </body>
        </html>
        """
        
        return html_content
